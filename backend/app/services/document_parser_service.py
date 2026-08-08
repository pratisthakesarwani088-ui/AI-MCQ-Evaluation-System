"""
document_parser_service.py
-----------------------------
Extracts plain text from an uploaded question-paper document
(PDF, DOCX, or TXT). This raw text is what gets stored on the
QuestionPaper record and, later, parsed into individual questions.
"""

import os


def extract_text_from_txt(path: str) -> str:
    """Reads a plain-text file, tolerating minor encoding issues."""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except OSError as exc:
        raise RuntimeError(f"Could not read the text file: {exc}")


def extract_text_from_pdf(path: str) -> str:
    """Extracts text from every page of a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "PDF support requires the 'pypdf' package. Please install it (see requirements.txt)."
        ) from exc

    try:
        reader = PdfReader(path)
    except Exception as exc:
        raise RuntimeError(
            f"Could not open the PDF file. It may be corrupted, password-protected, or not a valid PDF. Details: {exc}"
        )

    text_parts = []
    for page in reader.pages:
        try:
            text_parts.append(page.extract_text() or "")
        except Exception:
            # Skip a single unreadable page rather than failing the
            # whole document -- scanned/image-only pages sometimes
            # yield no extractable text.
            continue

    return "\n".join(text_parts)


def extract_text_from_docx(path: str) -> str:
    """Extracts text from every paragraph of a Word (.docx) document."""
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support requires the 'python-docx' package. Please install it (see requirements.txt)."
        ) from exc

    try:
        document = docx.Document(path)
    except Exception as exc:
        raise RuntimeError(
            f"Could not open the DOCX file. It may be corrupted or not a valid Word document. Details: {exc}"
        )

    paragraphs = [p.text for p in document.paragraphs]

    # Also pull text out of any tables in the document, since
    # questions/options are sometimes laid out in a table.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text_from_document(path: str, file_type: str) -> str:
    """
    Dispatches to the right extractor based on file_type ("PDF",
    "DOCX", or "TXT"), then validates the result isn't empty.

    Raises RuntimeError with a friendly message on any failure,
    including an unsupported file_type or an empty/unreadable
    document.
    """
    if not os.path.exists(path):
        raise RuntimeError("The uploaded file could not be found on disk.")

    file_type = (file_type or "").upper()

    if file_type == "PDF":
        text = extract_text_from_pdf(path)
    elif file_type == "DOCX":
        text = extract_text_from_docx(path)
    elif file_type == "TXT":
        text = extract_text_from_txt(path)
    else:
        raise RuntimeError(f"Unsupported document type: '{file_type}'. Only PDF, DOCX, and TXT are supported.")

    if not text or not text.strip():
        raise RuntimeError(
            "The document appears to be empty, or no text could be extracted from it "
            "(this can happen with scanned/image-only PDFs). Please try a different file "
            "or use Manual Question Entry instead."
        )

    return text.strip()
